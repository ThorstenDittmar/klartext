"""Tests for NarrativeImportService.

The NarrativeImportService is responsible for:
  1. Reading raw file content from disk (text formats).
  2. Handing that content to a NarrativeParser and wrapping the result
     in a Narrative domain object.
  3. For binary formats (e.g. DOCX) delegating file reading to a
     NarrativeFileParser registered for that extension.

Fake parsers below let us test the service in isolation without
depending on the correctness of any real parser.
"""

from pathlib import Path

import pytest

from api.exceptions.narrative import NarrativeFileNotFoundError, NarrativeParseError
from api.models.narrative import Scene
from api.parsers.narrative_file_parser import NarrativeFileParser
from api.parsers.narrative_parser import NarrativeParser
from api.services.narrative_import_service import NarrativeImportService

FIXTURE_PATH = (
    Path(__file__).parent
    / "fixtures"
    / "klartext-eine-geschichte-ueber-eine-geschichte"
    / "narrative.md"
)

NONEXISTENT_PATH = Path("/tmp/does_not_exist/narrative.md")


class FakeNarrativeParser(NarrativeParser):
    """Records the content it receives so tests can inspect it."""

    def __init__(self) -> None:
        self.received_content: str | None = None

    def parse(self, content: str) -> list[Scene]:
        self.received_content = content
        return [
            Scene.create(title="Szene 1", text="Fake text.", position=1),
            Scene.create(title="Szene 2", text="Auch fake.", position=2),
        ]


class FakeParserReturnsNothing(NarrativeParser):
    """Simulates a parser that finds no scenes – e.g. because the content is nonsense."""

    def parse(self, content: str) -> list[Scene]:
        return []


class FakeNarrativeFileParser(NarrativeFileParser):
    """Records the path it receives and returns two fixed scenes."""

    def __init__(self) -> None:
        self.received_path: Path | None = None

    def parse_file(self, path: Path) -> list[Scene]:
        self.received_path = path
        return [
            Scene.create(title="Szene 1", text="Erster Text.", position=1),
            Scene.create(title="Szene 2", text="Zweiter Text.", position=2),
        ]


class FakeFileParserReturnsNothing(NarrativeFileParser):
    """Simulates a file parser that finds no scenes."""

    def parse_file(self, path: Path) -> list[Scene]:
        return []


# --- Happy path ---


def test_narrative_import_service_passes_file_content_to_parser() -> None:
    """Expects the parser to receive the file content – not None and not empty."""
    parser = FakeNarrativeParser()
    service = NarrativeImportService(parser=parser)

    service.import_from_file(FIXTURE_PATH)

    assert parser.received_content is not None
    assert len(parser.received_content) > 0


def test_narrative_import_service_passes_actual_fixture_content_to_parser() -> None:
    """Expects the parser to receive the real file content, not a placeholder."""
    parser = FakeNarrativeParser()
    service = NarrativeImportService(parser=parser)

    service.import_from_file(FIXTURE_PATH)

    assert "Klartext" in parser.received_content  # type: ignore[operator]


def test_narrative_import_service_returns_narrative_with_scenes_from_parser() -> None:
    """Expects the returned Narrative to contain exactly the scenes the parser produced."""
    parser = FakeNarrativeParser()
    service = NarrativeImportService(parser=parser)

    narrative = service.import_from_file(FIXTURE_PATH)

    assert len(narrative.scenes) == 2


def test_narrative_import_service_with_real_parser_creates_scenes_from_fixture() -> None:
    """Expects the full chain (service + real parser + fixture) to produce non-empty scenes."""
    from api.parsers.markdown_narrative_parser import MarkdownNarrativeParser

    parser = MarkdownNarrativeParser()
    service = NarrativeImportService(parser=parser)

    narrative = service.import_from_file(FIXTURE_PATH)

    assert len(narrative.scenes) > 0
    assert all(scene.text for scene in narrative.scenes)


# --- Error cases ---


def test_narrative_import_service_raises_for_nonexistent_file() -> None:
    """Expects a NarrativeFileNotFoundError when the given path does not exist."""
    service = NarrativeImportService(parser=FakeNarrativeParser())

    with pytest.raises(NarrativeFileNotFoundError):
        service.import_from_file(NONEXISTENT_PATH)


def test_narrative_import_service_raises_for_wrong_filename_in_fixture_directory() -> None:
    """Expects a NarrativeFileNotFoundError when the directory exists but the file name is wrong."""
    wrong_path = FIXTURE_PATH.parent / "typo_narrative.md"
    service = NarrativeImportService(parser=FakeNarrativeParser())

    with pytest.raises(NarrativeFileNotFoundError):
        service.import_from_file(wrong_path)


def test_narrative_import_service_raises_for_empty_file(tmp_path: Path) -> None:
    """Expects a NarrativeParseError when the file exists but contains no content at all."""
    empty_file = tmp_path / "narrative.md"
    empty_file.write_text("")
    service = NarrativeImportService(parser=FakeNarrativeParser())

    with pytest.raises(NarrativeParseError):
        service.import_from_file(empty_file)


def test_narrative_import_service_raises_for_file_with_no_parseable_scenes(tmp_path: Path) -> None:
    """Expects a NarrativeParseError when the file has content but the parser finds no scenes."""
    nonsense_file = tmp_path / "narrative.md"
    nonsense_file.write_text("Das hier ist kein gültiges Narrativ-Format.")
    service = NarrativeImportService(parser=FakeParserReturnsNothing())

    with pytest.raises(NarrativeParseError):
        service.import_from_file(nonsense_file)


# ---------------------------------------------------------------------------
# DOCX / file-parser dispatch
# ---------------------------------------------------------------------------


def test_narrative_import_service_dispatches_docx_to_file_parser(tmp_path: Path) -> None:
    """Expects a .docx file to be handled by the registered file parser, not the text parser."""
    file_parser = FakeNarrativeFileParser()
    text_parser = FakeNarrativeParser()
    service = NarrativeImportService(
        parser=text_parser,
        file_parsers={".docx": file_parser},
    )
    # Create a minimal but real DOCX so the path exists
    from docx import Document as _Document

    doc_path = tmp_path / "narrative.docx"
    _Document().save(str(doc_path))

    service.import_from_file(doc_path)

    assert file_parser.received_path == doc_path
    assert text_parser.received_content is None


def test_narrative_import_service_uses_text_parser_for_non_docx(tmp_path: Path) -> None:
    """Expects a .md file to be handled by the text parser, not a file parser."""
    file_parser = FakeNarrativeFileParser()
    text_parser = FakeNarrativeParser()
    service = NarrativeImportService(
        parser=text_parser,
        file_parsers={".docx": file_parser},
    )
    md_path = tmp_path / "narrative.md"
    md_path.write_text("### Szene 1\n\nText.", encoding="utf-8")

    service.import_from_file(md_path)

    assert text_parser.received_content is not None
    assert file_parser.received_path is None


def test_narrative_import_service_builds_narrative_from_file_parser_result(
    tmp_path: Path,
) -> None:
    """Expects the Narrative returned by the service to contain the scenes from the file parser."""
    file_parser = FakeNarrativeFileParser()
    service = NarrativeImportService(
        parser=FakeNarrativeParser(),
        file_parsers={".docx": file_parser},
    )
    from docx import Document as _Document

    doc_path = tmp_path / "narrative.docx"
    _Document().save(str(doc_path))

    narrative = service.import_from_file(doc_path)

    assert len(narrative.scenes) == 2
    assert narrative.scenes[0].title == "Szene 1"
    assert narrative.scenes[1].title == "Szene 2"


def test_narrative_import_service_raises_parse_error_when_file_parser_returns_no_scenes(
    tmp_path: Path,
) -> None:
    """Expects a NarrativeParseError when the file parser finds no scenes in a DOCX."""
    service = NarrativeImportService(
        parser=FakeNarrativeParser(),
        file_parsers={".docx": FakeFileParserReturnsNothing()},
    )
    from docx import Document as _Document

    doc_path = tmp_path / "empty.docx"
    _Document().save(str(doc_path))

    with pytest.raises(NarrativeParseError):
        service.import_from_file(doc_path)


def test_narrative_import_service_raises_file_not_found_for_nonexistent_docx() -> None:
    """Expects a NarrativeFileNotFoundError when the .docx path does not exist."""
    service = NarrativeImportService(
        parser=FakeNarrativeParser(),
        file_parsers={".docx": FakeNarrativeFileParser()},
    )

    with pytest.raises(NarrativeFileNotFoundError):
        service.import_from_file(Path("/tmp/does_not_exist/narrative.docx"))


def test_narrative_import_service_derives_title_from_docx_filename(
    tmp_path: Path,
) -> None:
    """Expects the Narrative title to be derived from the DOCX filename stem."""
    service = NarrativeImportService(
        parser=FakeNarrativeParser(),
        file_parsers={".docx": FakeNarrativeFileParser()},
    )
    from docx import Document as _Document

    doc_path = tmp_path / "meine-geschichte.docx"
    _Document().save(str(doc_path))

    narrative = service.import_from_file(doc_path)

    assert narrative.title == "meine-geschichte"
