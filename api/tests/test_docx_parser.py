"""Tests for DocxNarrativeParser.

The parser receives a Path to a .docx file and returns Scene objects.
It knows nothing about services or databases – pure file transformation.

Convention: scene boundaries are plain-text paragraphs whose text
matches the pattern "Szene <number>" (e.g. "Szene 1", "Szene 2").
All paragraphs before the first scene marker are treated as preamble
and are ignored.  Empty paragraphs within a scene body are dropped.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from api.parsers.docx_narrative_parser import DocxNarrativeParser

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx(paragraphs: list[str], path: Path) -> Path:
    """Creates a DOCX file whose paragraphs are taken from *paragraphs*.

    Every string becomes one Normal-style paragraph.  Returns the path
    so callers can chain: ``parser.parse_file(_make_docx([...], tmp_path / "f.docx"))``.
    """
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Happy path
# ---------------------------------------------------------------------------


def test_docx_parser_returns_empty_list_for_document_without_scene_markers(
    tmp_path: Path,
) -> None:
    """Expects an empty list when no paragraph matches the 'Szene N' pattern."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Klartext", "Eine Geschichte.", "Kein Szenentrenner."], tmp_path / "t.docx"
    )

    scenes = parser.parse_file(doc_path)

    assert scenes == []


def test_docx_parser_returns_empty_list_for_document_with_no_paragraphs(
    tmp_path: Path,
) -> None:
    """Expects an empty list for an entirely blank DOCX – not an error."""
    parser = DocxNarrativeParser()
    doc = Document()
    doc_path = tmp_path / "empty.docx"
    doc.save(str(doc_path))

    scenes = parser.parse_file(doc_path)

    assert scenes == []


def test_docx_parser_creates_one_scene_from_single_scene_document(
    tmp_path: Path,
) -> None:
    """Expects exactly one Scene when the document contains one scene marker."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(["Szene 1", "Es war ein Dienstag."], tmp_path / "t.docx")

    scenes = parser.parse_file(doc_path)

    assert len(scenes) == 1


def test_docx_parser_creates_multiple_scenes(tmp_path: Path) -> None:
    """Expects one Scene per 'Szene N' marker found in the document."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Erster Text.", "Szene 2", "Zweiter Text.", "Szene 3", "Dritter Text."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert len(scenes) == 3


def test_docx_parser_assigns_title_from_scene_marker(tmp_path: Path) -> None:
    """Expects the scene title to be the full marker text (e.g. 'Szene 1')."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(["Szene 1", "Irgendein Text."], tmp_path / "t.docx")

    scenes = parser.parse_file(doc_path)

    assert scenes[0].title == "Szene 1"


def test_docx_parser_assigns_positions_starting_at_one(tmp_path: Path) -> None:
    """Expects positions to start at 1 and increase by 1 per scene."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Erster Text.", "Szene 2", "Zweiter Text."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert scenes[0].position == 1
    assert scenes[1].position == 2


def test_docx_parser_joins_body_paragraphs_with_double_newline(tmp_path: Path) -> None:
    """Expects multiple body paragraphs to be joined with a blank line between them."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Erster Absatz.", "Zweiter Absatz."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert "Erster Absatz." in scenes[0].text
    assert "Zweiter Absatz." in scenes[0].text
    assert "\n\n" in scenes[0].text


def test_docx_parser_ignores_paragraphs_before_first_scene_marker(
    tmp_path: Path,
) -> None:
    """Expects preamble paragraphs (before 'Szene 1') to be discarded."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Klartext", "Eine Geschichte über eine Geschichte", "Szene 1", "Der eigentliche Text."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert len(scenes) == 1
    assert "Klartext" not in scenes[0].text
    assert "Eine Geschichte" not in scenes[0].text


def test_docx_parser_skips_scene_with_no_body_paragraphs(tmp_path: Path) -> None:
    """Expects a scene marker immediately followed by another marker to be skipped."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Szene 2", "Zweiter Text."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert len(scenes) == 1
    assert scenes[0].title == "Szene 2"


def test_docx_parser_strips_whitespace_from_scene_text(tmp_path: Path) -> None:
    """Expects no leading or trailing whitespace in the final scene text."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(["Szene 1", "  Text mit Leerzeichen.  "], tmp_path / "t.docx")

    scenes = parser.parse_file(doc_path)

    assert scenes[0].text == scenes[0].text.strip()


def test_docx_parser_drops_empty_paragraphs_within_scene_body(tmp_path: Path) -> None:
    """Expects blank paragraphs inside a scene to be silently dropped."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Erster Satz.", "", "Zweiter Satz."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    assert "" not in scenes[0].text.split("\n\n")


def test_docx_parser_handles_scene_marker_with_extra_whitespace(tmp_path: Path) -> None:
    """Expects 'Szene  2' (multiple spaces) to be recognised as a scene marker."""
    parser = DocxNarrativeParser()
    doc_path = _make_docx(["Szene  2", "Text."], tmp_path / "t.docx")

    scenes = parser.parse_file(doc_path)

    assert len(scenes) == 1


def test_docx_parser_does_not_treat_partial_match_as_scene_marker(
    tmp_path: Path,
) -> None:
    """Expects 'Szene 1 – Der Abend' not to be treated as a pure scene marker.

    Only paragraphs whose stripped text is exactly 'Szene <number>' are markers.
    The full title (with em-dash suffix) belongs to the body.
    """
    parser = DocxNarrativeParser()
    doc_path = _make_docx(
        ["Szene 1", "Szene 1 – Der Abend", "Eigentlicher Fließtext."],
        tmp_path / "t.docx",
    )

    scenes = parser.parse_file(doc_path)

    # Only one scene (from "Szene 1"); the long line is body text
    assert len(scenes) == 1
    assert "Szene 1 – Der Abend" in scenes[0].text


# ---------------------------------------------------------------------------
# Error cases
# ---------------------------------------------------------------------------


def test_docx_parser_raises_for_nonexistent_file() -> None:
    """Expects a FileNotFoundError when the given path does not exist."""
    parser = DocxNarrativeParser()

    with pytest.raises(FileNotFoundError):
        parser.parse_file(Path("/tmp/does_not_exist/narrative.docx"))
